# -*- coding: utf-8 -*-
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "bridge-inspection-system"
DOCS = PROJECT / "docs"
ASSETS = DOCS / "overview-assets"
TEMPLATE = ROOT / "4. 概要设计示例.docx"
OUTPUT = DOCS / "公路桥梁初始检查信息系统概要设计.docx"
FALLBACK_OUTPUT = DOCS / "公路桥梁初始检查信息系统概要设计_PowerDesigner模型版.docx"
MARKDOWN = DOCS / "公路桥梁初始检查信息系统概要设计.md"

FONT = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")


def font(size, bold=False):
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT), size)


def rounded(draw, box, fill, outline="#CBD5E1", width=2, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, text_font, fill="#172033", spacing=8):
    x1, y1, x2, y2 = box
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, width=max(4, int((x2 - x1) / (text_font.size * 1.05)))))
    heights = [draw.textbbox((0, 0), line, font=text_font)[3] for line in lines]
    total = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, line_height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=text_font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=text_font, fill=fill)
        y += line_height + spacing


def arrow(draw, start, end, color="#64748B", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    left = (ex - ux * 18 - uy * 8, ey - uy * 18 + ux * 8)
    right = (ex - ux * 18 + uy * 8, ey - uy * 18 - ux * 8)
    draw.polygon([end, left, right], fill=color)


def create_function_diagram(path):
    image = Image.new("RGB", (1800, 1180), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 38), "公路桥梁初始检查信息系统总体功能架构", font=font(42, True), fill="#0F3D3A")
    center = (590, 115, 1210, 265)
    rounded(draw, center, "#0F766E", "#0F766E", 3, 24)
    center_text(draw, center, "公路桥梁初始检查信息系统", font(34, True), "white", 12)
    modules = [
        ("1 身份认证与角色门户", "登录、注册、修改密码\n五类角色独立工作台", "#ECFDF5"),
        ("2 基础数据与矩阵", "路线、桥型、部位、部件\n桥型部件矩阵、初检项目矩阵", "#EFF6FF"),
        ("3 桥梁档案与A卡", "桥梁总体信息、照片、档案\nD-1至D-N具体部件实例", "#FFF7ED"),
        ("4 初始检查与B表", "每桥一份有效记录\n公共检测项与桥墩分类记录", "#F5F3FF"),
        ("5 定期检查与C表", "一桥多次定检\n按桥型生成部件检查与病害", "#FEF2F2"),
        ("6 任务审核与报告", "任务分派、接收、完成\n审核归档、退回、报告生成", "#F0FDFA"),
        ("7 地图查询与综合检索", "重庆桥梁地图标注\n桥梁、检查、病害和报告查询", "#F0F9FF"),
        ("8 系统管理", "用户、角色、权限\n日志、备份和文件管理", "#F8FAFC"),
    ]
    box_w, box_h = 690, 175
    xs = [105, 1005]
    ys = [355, 565, 775, 985]
    left_rail, right_rail = 65, 1735
    rail_top, rail_bottom = 315, 1072
    draw.line([(center[0] + 110, center[3]), (left_rail, center[3]), (left_rail, rail_bottom)], fill="#94A3B8", width=4)
    draw.line([(center[2] - 110, center[3]), (right_rail, center[3]), (right_rail, rail_bottom)], fill="#94A3B8", width=4)
    for index, (title, desc, fill) in enumerate(modules):
        x, y = xs[index % 2], ys[index // 2]
        box = (x, y, x + box_w, y + box_h)
        rounded(draw, box, fill, "#94A3B8", 2, 18)
        draw.text((x + 26, y + 22), title, font=font(27, True), fill="#0F3D3A")
        draw.multiline_text((x + 26, y + 78), desc, font=font(22), fill="#334155", spacing=12)
        if index % 2 == 0:
            arrow(draw, (left_rail, y + box_h // 2), (x - 10, y + box_h // 2), "#94A3B8", 3)
        else:
            arrow(draw, (right_rail, y + box_h // 2), (x + box_w + 10, y + box_h // 2), "#94A3B8", 3)
    image.save(path)


def create_architecture_diagram(path):
    image = Image.new("RGB", (1800, 1260), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 38), "系统前端、后端与数据库技术架构", font=font(42, True), fill="#0F3D3A")
    layers = [
        (95, 145, 1705, 365, "用户与表示层", "浏览器 / Vue 3 / Element Plus / Vue Router / Pinia\n角色门户、资源管理页、A/B/C卡片、重庆桥梁地图、报告与任务页面", "#ECFDF5"),
        (95, 435, 1705, 675, "接口与业务层", "Spring Boot 3.3.5 / REST API / Spring Security / JWT\n认证注册、通用CRUD、桥梁聚合详情、任务流转、矩阵生成、报告、文件服务", "#EFF6FF"),
        (95, 745, 1180, 1035, "数据持久层", "JdbcTemplate / HikariCP / Flyway\nMySQL 8.0、InnoDB、utf8mb4、外键、唯一约束、索引、中文查询视图", "#FFF7ED"),
        (1240, 745, 1705, 1035, "文件与外部服务", "uploads文件目录\n桥梁照片与档案附件\n高德地图Web端服务", "#F5F3FF"),
    ]
    for x1, y1, x2, y2, title, desc, fill in layers:
        rounded(draw, (x1, y1, x2, y2), fill, "#94A3B8", 3, 24)
        draw.text((x1 + 34, y1 + 28), title, font=font(30, True), fill="#0F3D3A")
        draw.multiline_text((x1 + 34, y1 + 90), desc, font=font(24), fill="#334155", spacing=14)
    arrow(draw, (900, 365), (900, 425), "#0F766E", 5)
    arrow(draw, (700, 675), (700, 735), "#0F766E", 5)
    arrow(draw, (1360, 675), (1460, 735), "#0F766E", 5)
    draw.text((705, 387), "HTTPS / JSON", font=font(20, True), fill="#0F766E")
    draw.text((720, 698), "SQL", font=font(20, True), fill="#0F766E")
    draw.text((1380, 690), "文件与地图调用", font=font(20, True), fill="#0F766E")
    rounded(draw, (95, 1110, 1705, 1200), "#0F766E", "#0F766E", 2, 18)
    center_text(draw, (95, 1110, 1705, 1200), "部署：前端5173端口　后端8080端口　MySQL 3306端口　B/S无状态访问", font(24, True), "white")
    image.save(path)


def create_data_diagram(path):
    image = Image.new("RGB", (1900, 1350), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 38), "概念数据模型分区图（按业务域拆分）", font=font(42, True), fill="#0F3D3A")
    groups = [
        (80, 135, 600, 450, "基础与桥梁结构域", ["路线 → 桥梁", "桥梁类型 → 桥型部件配置", "部位 → 部件", "桥梁 → 桥梁具体部件", "D-1桥墩/桥桩/基础，D-2至D-N其他部件"], "#ECFDF5"),
        (690, 135, 1210, 450, "桥梁档案域（A卡）", ["桥梁基本状况卡片", "桥梁档案资料项", "桥梁档案资料记录", "检测评定历史", "照片/档案文件"], "#FFF7ED"),
        (1300, 135, 1820, 450, "用户权限域", ["角色 → 用户", "用户创建业务记录", "操作日志", "数据库备份记录", "JWT身份与权限集合"], "#F5F3FF"),
        (80, 560, 600, 1010, "初始检查域（B表）", ["桥梁 → 初始检查记录", "每座桥梁一份有效B表", "初始检查项目定义", "桥型初检项目配置", "初始检查检测项目", "初始检查具体部件检测记录", "病害/缺损 → 附件"], "#EFF6FF"),
        (690, 560, 1210, 1010, "定期检查域（C表）", ["桥梁 → 多次定期检查记录", "定期检查 → 部件检查记录", "部件检查关联桥梁具体部件", "缺损程度与技术状况等级", "病害/缺损 → 附件", "按桥型模板生成检查行"], "#FEF2F2"),
        (1300, 560, 1820, 1010, "任务与报告域", ["检查任务", "任务分配", "任务状态历史", "检查报告", "任务→检查→审核→归档", "报告文件下载"], "#F0FDFA"),
    ]
    for x1, y1, x2, y2, title, items, fill in groups:
        rounded(draw, (x1, y1, x2, y2), fill, "#94A3B8", 3, 22)
        draw.text((x1 + 28, y1 + 24), title, font=font(27, True), fill="#0F3D3A")
        y = y1 + 82
        for item in items:
            draw.ellipse((x1 + 30, y + 8, x1 + 42, y + 20), fill="#0F766E")
            draw.text((x1 + 57, y), item, font=font(20), fill="#334155")
            y += 48
    arrow(draw, (600, 290), (680, 290), "#0F766E", 4)
    arrow(draw, (955, 450), (955, 550), "#0F766E", 4)
    arrow(draw, (600, 780), (680, 780), "#0F766E", 4)
    arrow(draw, (1210, 780), (1290, 780), "#0F766E", 4)
    rounded(draw, (250, 1110, 1650, 1260), "#0F766E", "#0F766E", 2, 22)
    center_text(draw, (250, 1110, 1650, 1260), "核心链：路线 → 桥梁 → 具体部件 → A/B/C卡片 → 部件检测 → 病害 → 附件\n模型以业务域拆分展示，PowerDesigner中使用多个子图避免关系线形成蜘蛛网", font(25, True), "white", 12)
    image.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, east_asia="宋体", size=10.5, bold=None, color=None):
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5, bold)
    return paragraph


def add_label(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 10.5, True)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run("• " + item)
        set_run_font(run, "宋体", 10.5)


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 18
    elif level == 2:
        size = 14
    else:
        size = 12
    for run in paragraph.runs:
        set_run_font(run, "黑体", size, True)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5)


def add_picture(doc, path, width=16.0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width))


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "CBD5E1")
        borders.append(border)
    table_properties.append(borders)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, title in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, "0F766E")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        set_run_font(run, "黑体", 9.5, True, "FFFFFF")
        if widths:
            cell.width = Cm(widths[index])
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, "宋体", 9)
            if widths:
                cell.width = Cm(widths[index])
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    header = section.header.paragraphs[0]
    header.text = "《程序设计综合实践II》概要设计"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.runs[0], "宋体", 9)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def generate_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    DOCS.mkdir(parents=True, exist_ok=True)
    function_diagram = ASSETS / "功能架构图.png"
    architecture_diagram = ASSETS / "技术架构图.png"
    data_diagram = ASSETS / "概念数据模型分区图.png"
    pd_diagram_dir = ASSETS / "powerdesigner"
    pd_cdm_diagram = pd_diagram_dir / "PowerDesigner_CDM_基础字典与桥型矩阵.png"
    pd_ldm_diagram = pd_diagram_dir / "PowerDesigner_LDM_初始检查B表.png"
    pd_pdm_diagram = pd_diagram_dir / "PowerDesigner_PDM_定期检查C表.png"
    create_function_diagram(function_diagram)
    create_architecture_diagram(architecture_diagram)
    create_data_diagram(data_diagram)

    doc = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_body(doc)
    configure_document(doc)

    add_heading(doc, "第3部分 系统设计", 1)
    add_body(doc, "以前期需求分析、JTG 5120—2021《公路桥涵养护规范》、桥梁模型示例图、系统总体功能设计图以及当前已实现的前端、后端和数据库为基础，设计《公路桥梁初始检查信息系统》的总体功能、软件架构和数据结构。本系统面向重庆及周边桥梁建档、初始检查、定期检查、病害记录、审核归档和地图查询，不建设独立的养护处治业务子系统。")

    add_heading(doc, "3.1 系统功能设计", 2)
    add_body(doc, "系统采用按角色授权的功能组织方式，将用户需求转换为身份认证、基础数据、桥梁档案、初始检查、定期检查、任务审核、查询统计和系统管理八类功能。各功能通过统一REST接口访问业务数据库和文件存储，系统总体功能架构如图3.1所示。")
    add_picture(doc, function_diagram, 16.0)
    add_caption(doc, "图3.1 系统总体功能架构")

    modules = [
        ("3.1.1 身份认证与角色门户", "完成用户登录、公开注册、身份校验、修改密码、退出登录以及不同角色首页跳转。系统提供系统管理员、桥梁工程师、检查人员、审核人员和查询人员五种角色门户。", ["注册用户默认写入用户表并授予viewer查询人员角色。", "登录成功后后端签发JWT，前端保存令牌及用户权限。", "路由守卫和后端安全过滤器共同限制越权访问。"]),
        ("3.1.2 基础数据与桥型矩阵", "维护路线、桥梁类型、部位、部件、检查项目定义以及桥型配置关系。部位与部件为一对多关系，桥型通过配置表确定应包含和检查的部件。", ["支持6类桥型和151条桥型—部位—部件配置关系。", "支持47个初检项目和220条桥型—初检项目适用关系。", "新建桥梁或检查记录时可根据矩阵自动生成实际部件或检测项目行。"]),
        ("3.1.3 桥梁档案与基本状况卡片（A表）", "管理桥梁编号、路线、桥型、行政区划、坐标、技术指标、档案资料、评定历史、照片和实际部件。每座桥梁对应一份A表。", ["桥梁详情页显示桥梁图片和总体信息，并提供A、B、C三类卡片切换。", "结构信息不采用横向动态列，而按D-1至D-N分区展示。", "D-1固定为桥墩、桥桩与基础实例，每个实际实例占一条记录；其余部件按类别形成D-2、D-3等分区。", "桥梁照片支持上传，首张图片可作为详情封面。"]),
        ("3.1.4 初始检查记录（B表）", "管理初始检查基本信息、规范检测项目、具体桥墩或部件检测结果、病害和附件。每座桥梁只保留一份有效初始检查记录。", ["公共检测项目依据桥型初检项目矩阵生成。", "桥墩、桥桩等具体部件的检测结果按实际实例分类保存。", "归档记录保存桥梁关键字段快照，避免基础档案修改影响历史结果。"]),
        ("3.1.5 定期检查记录（C表）", "管理同一桥梁历次定期检查、技术状况等级、部件检查、评分、病害、照片和下次检查建议。一座桥梁可以对应多份C表。", ["创建定检时依据桥型部件配置自动生成检查行。", "每次定检可分别记录各桥墩、桥桩和其他具体部件的检查结果。", "支持按时间查看同桥梁多次定检记录。"]),
        ("3.1.6 检查任务、审核与报告", "提供任务创建、分配、接收、完成、审核、退回、取消和状态历史记录，并根据任务生成检查报告。", ["任务状态流转为待分配、已分配、进行中、待审核、已完成或退回。", "审核人员拥有独立审核门户，不与录入门户混用。", "报告文件通过受控接口生成和下载。"]),
        ("3.1.7 重庆桥梁地图与综合查询", "通过高德地图Web端服务展示重庆周边桥梁坐标标注。点击标注显示桥梁名称，点击名称进入桥梁详情和A/B/C卡片。", ["支持按桥梁、路线、桥型、检查日期、技术等级和病害条件查询。", "查询人员只读访问桥梁、检查、病害和报告数据。", "地图密钥由环境变量或后端配置提供，前端不保存Web服务安全密钥。"]),
        ("3.1.8 系统管理", "由系统管理员维护用户、角色与权限、操作日志、数据库备份和系统版本信息。", ["密码采用BCrypt散列存储。", "用户状态支持启用和停用。", "关键操作写入操作日志，备份记录保存文件路径和执行状态。"]),
    ]
    for heading, description, rules in modules:
        add_heading(doc, heading, 3)
        add_label(doc, "功能描述")
        add_body(doc, description)
        add_label(doc, "操作数据描述")
        add_bullets(doc, rules)

    add_heading(doc, "3.1.9 角色与功能权限设计", 3)
    add_table(doc, ["角色", "主要职责", "可访问功能"], [
        ["系统管理员", "系统配置和全局管理", "全部角色门户、基础数据、A/B/C卡片、审核、查询、用户、角色、日志、备份"],
        ["桥梁工程师", "桥梁建档和专业数据维护", "基础数据、桥梁档案、初始检查、定期检查、查询统计"],
        ["检查人员", "现场检查数据填报", "检查任务、初始检查、定期检查、病害、报告和查询"],
        ["审核人员", "检查结果审核与归档", "初检审核、定检审核、报告审核和只读查询"],
        ["查询人员", "业务数据只读查询", "重庆桥梁地图、桥梁、检查、病害和报告查询"],
    ], [2.4, 4.0, 9.0])

    add_heading(doc, "3.2 系统总体架构设计", 2)
    add_body(doc, "系统采用B/S前后端分离架构。浏览器负责界面展示和交互，Spring Boot服务负责认证、权限和业务处理，MySQL保存结构化数据，uploads目录保存桥梁照片、档案和报告文件，高德地图提供桥梁空间展示能力。系统技术架构如图3.2所示。")
    add_picture(doc, architecture_diagram, 16.0)
    add_caption(doc, "图3.2 系统前端、后端与数据库技术架构")

    add_heading(doc, "3.2.1 前端设计", 3)
    add_body(doc, "前端使用Vue 3、Vite、Element Plus、Vue Router、Pinia、Axios和Lucide图标库。页面采用路由级功能拆分与通用资源页面相结合的设计。")
    add_table(doc, ["前端模块", "主要页面或组件", "设计说明"], [
        ["认证模块", "LoginView、Pinia认证状态", "提供登录、注册、令牌保存、退出和角色首页跳转"],
        ["角色工作台", "Admin/Engineer/Inspector/Reviewer/Viewer Dashboard", "不同角色进入不同首页并显示不同功能入口"],
        ["资源管理", "ResourceView、MatrixView", "以资源编码配置路线、桥型、部位、部件、矩阵和业务记录的表格及表单"],
        ["桥梁详情", "BridgeProfileView", "显示桥梁封面、总体信息以及A/B/C三类卡片"],
        ["动态部件分区", "BridgeDynamicMatrix", "按实际部件类别生成D-1至D-N，每个实例一行"],
        ["检查卡片", "BridgeInitialCard、BridgePeriodicCards", "B表每桥一份，C表一桥多份并按具体部件分类"],
        ["地图查询", "BridgeMapView", "加载高德地图标注并跳转桥梁详情"],
        ["任务与报告", "TaskView、ReportView", "完成任务流转、审核和报告生成下载"],
    ], [3.0, 4.6, 7.8])
    add_body(doc, "标签交互采用当前路由和组件独立状态控制。未选中的标签使用深色文字和浅色背景；选中标签反转为深色背景、浅色文字，且每个菜单和卡片使用唯一键值，避免点击一个标签导致其他标签同时变亮。")

    add_heading(doc, "3.2.2 后端设计", 3)
    add_body(doc, "后端使用Java 17、Spring Boot 3.3.5、Spring Security、JWT、JdbcTemplate、HikariCP和Flyway。采用Controller—Service—Database三层结构，返回统一JSON响应。")
    add_table(doc, ["后端模块", "主要接口", "主要职责"], [
        ["认证服务", "/api/auth", "登录、注册、获取当前用户、修改密码、JWT签发和BCrypt校验"],
        ["通用资源服务", "/api/{resource}", "根据资源白名单执行分页查询、详情、新增、修改和删除"],
        ["桥梁聚合服务", "/api/bridge-profiles", "聚合桥梁、具体部件、A/B/C记录、档案、照片和地图点位"],
        ["矩阵服务", "/api/matrices", "查询桥型矩阵并批量生成桥梁部件和初检项目"],
        ["任务服务", "/api/tasks", "任务分配、接收、完成、审核、退回、取消及状态历史"],
        ["报告服务", "/api/reports", "生成检查报告并提供下载"],
        ["文件服务", "/api/files", "上传、在线查看和下载照片、档案及报告文件"],
        ["仪表盘服务", "/api/dashboard", "按角色返回工作台统计信息"],
    ], [3.0, 4.4, 8.0])

    add_heading(doc, "3.2.3 接口与数据流设计", 3)
    add_bullets(doc, [
        "浏览器通过Axios调用/api路径，开发环境由Vite代理到8080端口。",
        "除登录、注册、健康检查和文件预览外，接口均要求携带JWT。",
        "桥梁详情采用聚合接口一次返回桥梁、部件、档案、初检、定检和附件，减少页面多次请求。",
        "文件上传使用multipart/form-data，数据库只保存文件元数据和相对路径。",
        "高德Web端Key用于地图显示；Web服务安全密钥只允许在后端环境变量中配置。",
    ])

    add_heading(doc, "3.3 数据模型设计", 2)
    add_body(doc, "概念数据模型、逻辑数据模型和物理数据模型均使用PowerDesigner 16.5实际生成。三份正式模型均依据当前运行数据库bridge_inspection_v2的元数据建立，包含30个业务对象、294个字段和56条关系，并分别设置6个业务子图。文档中的模型图片由PowerDesigner直接导出，正式设计成果以.cdm、.ldm和.pdm文件为准。")
    add_table(doc, ["模型类型", "PowerDesigner正式文件", "规模"], [
        ["概念数据模型CDM", "powerdesigner-mysql-er/公路桥梁初始检查信息系统概念模型_当前版.cdm", "30个实体、294个属性、56条关系、6个概念子图"],
        ["逻辑数据模型LDM", "powerdesigner-mysql-er/公路桥梁初始检查信息系统逻辑模型_当前版.ldm", "30个实体、294个属性、56条关系、6个逻辑子图"],
        ["物理数据模型PDM", "powerdesigner-mysql-er/公路桥梁初始检查信息系统物理模型_当前版.pdm", "30张表、294个字段、56条外键、68个索引、6个物理子图"],
    ], [3.2, 8.2, 5.0])
    add_heading(doc, "3.3.1 概念数据模型", 3)
    add_body(doc, "概念模型围绕桥梁档案、初始检查和定期检查三类核心业务组织，并拆分基础结构、用户权限、任务报告和附件等辅助业务域。为提高PowerDesigner模型可读性，CDM建立“核心业务总览、基础字典与桥型矩阵、桥梁档案与A卡、初始检查B表、定期检查C表、任务报告与用户权限”6个概念子图。图3.3为PowerDesigner直接导出的基础字典与桥型矩阵子图。")
    add_picture(doc, pd_cdm_diagram, 16.0)
    add_caption(doc, "图3.3 PowerDesigner概念数据模型——基础字典与桥型矩阵")
    add_body(doc, "概念模型的关键抽象是“标准部件定义—桥型部件配置—桥梁具体部件实例”。桥墩、桥桩、基础、支座、主梁、斜拉索等实际对象均作为桥梁具体部件记录，不使用固定数量的数据库列。A表按D-1至D-N显示这些实例；B表和C表通过桥梁具体部件编号保存每个实际桥墩或部件的检测结果。")

    add_heading(doc, "3.3.2 核心实体关系", 3)
    add_table(doc, ["关系", "基数", "设计说明"], [
        ["路线—桥梁", "1:N", "一条路线包含多座桥梁，桥梁保存路线编号外键"],
        ["桥梁类型—桥型部件配置", "1:N", "定义每种桥型应包含和检查的部位、部件"],
        ["部位—部件", "1:N", "部件必须归属一个部位，防止部位与部件关系混乱"],
        ["桥梁—桥梁具体部件", "1:N", "一座桥梁包含多个实际桥墩、桥桩、基础和其他部件"],
        ["桥梁—初始检查记录", "1:1有效", "每座桥梁只允许一份有效B表"],
        ["桥梁—定期检查记录", "1:N", "每座桥梁可保存多次C表"],
        ["初始检查—具体部件检测", "1:N", "按实际桥墩或部件保存初检结果"],
        ["定期检查—部件检查记录", "1:N", "每次定检包含多个实际部件检查结果"],
        ["部件检查—病害/缺损", "1:N", "一条部件检查可记录多个缺损"],
        ["业务对象—照片/档案文件", "1:N", "桥梁、档案、检查、部件检查或病害可关联多个附件"],
        ["角色—用户", "1:N", "一个角色包含多个用户，用户登录后获得角色权限"],
        ["检查任务—任务状态历史", "1:N", "完整记录任务流转过程"],
    ], [5.0, 2.0, 8.4])

    add_heading(doc, "3.3.3 逻辑模型设计", 3)
    add_body(doc, "逻辑模型由PowerDesigner LDM生成，使用实体属性、主标识、非标识关系和既有外键属性表达概念实体向关系模式的转换。逻辑模型包含30个实体、294个属性和56条关系，与当前数据库业务字段一致，不包含PowerDesigner自动重复迁移的外键属性。V6迁移已删除养护处治表和处治类别表，因此逻辑模型不再包含独立养护处治实体。")
    add_table(doc, ["业务域", "主要逻辑表", "说明"], [
        ["用户权限", "tb_role、tb_user", "角色、权限集合、用户账号、密码和状态"],
        ["基础结构", "tb_route、tb_bridge_type、tb_part、tb_component、tb_bridge_type_component_config", "路线、桥型、部位、部件和桥型模板"],
        ["桥梁档案", "tb_bridge、tb_bridge_specific_component、tb_archive_item、tb_bridge_archive_record、tb_evaluation_history、tb_attachment", "A表、实际部件、档案、评定和附件"],
        ["初始检查", "tb_initial_inspection_item_definition、tb_bridge_type_initial_item_config、tb_initial_inspection、tb_initial_inspection_item、tb_initial_component_inspection", "B表、项目模板、公共项目和具体部件检测"],
        ["定期检查", "tb_periodic_inspection、tb_component_inspection、tb_defect、tb_defect_degree、tb_rating_level", "C表、部件检查、病害和等级字典"],
        ["任务与报告", "tb_inspection_task、tb_task_assignment、tb_task_status_history、tb_report", "任务生命周期和报告"],
        ["系统运行", "tb_operation_log、tb_backup_record、tb_check_category", "日志、备份和检测类别字典"],
    ], [3.0, 7.6, 7.2])
    add_picture(doc, pd_ldm_diagram, 16.0)
    add_caption(doc, "图3.4 PowerDesigner逻辑数据模型——初始检查B表")

    add_heading(doc, "3.3.4 物理模型设计", 3)
    add_bullets(doc, [
        "数据库平台为MySQL 8.0，存储引擎使用InnoDB，字符集使用utf8mb4。",
        "数据库名默认为bridge_inspection_v2，连接池使用HikariCP。",
        "主键根据业务稳定性选择编码型VARCHAR或自增INT；业务编码设置唯一约束。",
        "外键用于保证路线、桥型、部位、部件、检查和附件的引用完整性。",
        "初始检查表对bridge_code设置唯一约束，保证每桥一份有效B表。",
        "初始具体部件检测表对初检编号、具体部件编号和项目编码设置组合唯一约束。",
        "桥梁照片、档案和报告文件保存在uploads目录，数据库保存名称、类型、大小、路径、上传人和时间。",
        "Flyway按V1至V13执行建表、字典数据、示例数据、中文视图、地图坐标和桥墩实例迁移。",
        "底层物理表使用稳定英文标识；同时建立中文数据库视图，供PowerDesigner展示、教学检查和中文查询使用。",
    ])
    add_picture(doc, pd_pdm_diagram, 16.0)
    add_caption(doc, "图3.5 PowerDesigner物理数据模型——定期检查C表")

    add_heading(doc, "3.3.5 关键数据约束", 3)
    add_table(doc, ["约束项", "约束内容"], [
        ["桥梁编号", "全局唯一，如G75-001，并作为A/B/C关联桥梁的业务编码"],
        ["部位与部件", "每个部件必须引用一个部位；桥型配置同时引用桥型、部位和部件"],
        ["具体部件实例", "每个实际桥墩、桥桩、基础、索、支座等一条记录，使用部件序号区分"],
        ["A表", "每座桥梁一份，由桥梁主表及其档案、评定、附件和具体部件聚合形成"],
        ["B表", "每座桥梁一份有效初始检查；公共项目与具体部件项目分表保存"],
        ["C表", "一座桥梁可有多份定期检查，每份包含多个部件检查和病害记录"],
        ["状态流转", "检查记录采用草稿、待审核、已归档、退回等状态；任务保存完整状态历史"],
        ["删除策略", "桥梁和业务数据优先采用状态字段逻辑删除，避免破坏历史记录"],
        ["附件关联", "附件上传前业务对象必须存在，文件元数据与物理文件路径保持一致"],
    ], [4.0, 11.6])

    add_heading(doc, "3.4 前端页面概要设计", 2)
    add_table(doc, ["页面", "访问角色", "主要内容"], [
        ["登录/注册页", "公开访问", "登录、查询人员注册、密码规则提示"],
        ["角色工作台", "对应角色", "角色统计、待办事项和快捷入口"],
        ["基础数据页", "管理员、桥梁工程师", "路线、桥型、部位、部件和两类矩阵维护"],
        ["桥梁档案页", "管理员、桥梁工程师", "桥梁列表、A表、具体部件和档案资料"],
        ["桥梁详情页", "全部登录角色", "桥梁图片、总体信息、A/B/C卡片切换"],
        ["初始检查页", "管理员、工程师、检查人员", "任务、B表、检测项目、具体桥墩检测、病害和报告"],
        ["定期检查页", "管理员、工程师、检查人员", "任务、多份C表、部件检查、病害和报告"],
        ["审核归档页", "管理员、审核人员", "初检、定检和报告审核"],
        ["重庆桥梁地图", "全部登录角色", "高德地图标注、桥名信息窗和详情跳转"],
        ["系统管理页", "管理员", "用户、角色权限、日志和备份"],
    ], [3.5, 3.3, 8.8])

    add_heading(doc, "3.5 安全与运行设计", 2)
    add_bullets(doc, [
        "认证采用无状态JWT，后端在请求过滤阶段解析用户、角色和权限。",
        "密码使用BCrypt，不保存明文密码；注册账号默认权限最小化。",
        "前端路由权限用于改善交互，后端接口认证用于最终安全控制。",
        "CORS仅允许本机开发地址，生产环境应改为正式域名白名单。",
        "上传文件限制大小和类型，存储文件名使用随机名称，避免覆盖和路径注入。",
        "数据库账号、JWT密钥、高德Key和安全密钥应通过环境变量配置，不写入公开前端代码。",
        "系统运行端口为前端5173、后端8080、MySQL 3306；生产环境可由Nginx统一反向代理。",
    ])

    add_heading(doc, "3.6 部署与开发环境", 2)
    add_table(doc, ["类别", "配置"], [
        ["前端", "Node.js、Vue 3.5、Vite 5.4、Element Plus 2.8，使用npm构建"],
        ["后端", "JDK 17、Spring Boot 3.3.5、Maven，打包为可执行JAR"],
        ["数据库", "MySQL 8.0，Flyway自动执行db/v2目录迁移"],
        ["文件", "项目uploads目录，按业务和日期组织桥梁照片、档案及报告"],
        ["开发地址", "前端http://localhost:5173，后端http://localhost:8080"],
        ["生产建议", "Nginx + 前端静态资源 + Spring Boot服务 + MySQL + 独立文件目录"],
    ], [3.6, 12.0])

    try:
        doc.save(str(OUTPUT))
    except PermissionError:
        doc.save(str(FALLBACK_OUTPUT))
    return function_diagram, architecture_diagram, data_diagram


def generate_markdown():
    content = """# 第3部分 系统设计

本概要设计以需求分析、JTG 5120—2021、桥梁模型示例图、系统总体功能设计图及当前实现为依据。系统面向重庆周边桥梁建档、初始检查、定期检查、病害、审核、报告和地图查询，不包含独立养护处治子系统。

## 3.1 系统功能设计

系统包含身份认证与角色门户、基础数据与矩阵、桥梁档案与A卡、初始检查B表、定期检查C表、任务审核与报告、重庆桥梁地图与综合查询、系统管理八类功能。

## 3.2 系统总体架构设计

- 前端：Vue 3、Vite、Element Plus、Vue Router、Pinia、Axios。
- 后端：Java 17、Spring Boot 3.3.5、Spring Security、JWT、JdbcTemplate、Flyway。
- 数据库：MySQL 8.0、InnoDB、utf8mb4、中文查询视图。
- 文件与地图：uploads目录、高德地图Web端服务。

## 3.3 数据模型设计

核心关系为：路线 → 桥梁 → 桥梁具体部件；桥梁 → A表；桥梁 → 一份有效B表；桥梁 → 多份C表；检查 → 部件检查 → 病害 → 附件。桥墩、桥桩与基础固定显示为D-1，每个实际实例一条记录，其余部件依次形成D-2至D-N。

完整章节、角色权限表、前后端模块表、逻辑表清单、约束和部署说明见同目录Word文档《公路桥梁初始检查信息系统概要设计.docx》。
"""
    MARKDOWN.write_text(content, encoding="utf-8")


if __name__ == "__main__":
    diagrams = generate_doc()
    generate_markdown()
    generated_output = FALLBACK_OUTPUT if FALLBACK_OUTPUT.exists() and FALLBACK_OUTPUT.stat().st_mtime >= OUTPUT.stat().st_mtime else OUTPUT
    print(generated_output)
    print(MARKDOWN)
    for diagram in diagrams:
        print(diagram)
