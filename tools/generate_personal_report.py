from __future__ import annotations

import html
import math
import os
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from matplotlib.patches import Ellipse, FancyArrowPatch, FancyBboxPatch, Polygon
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[2]
SOURCE_MD = WORKSPACE / "成员设计报告_肖皓轩_632407070128.md"
TARGET_DOCX = next(WORKSPACE.glob("4*.docx"))
ASSET_DIR = WORKSPACE / "bridge-inspection-system" / "docs" / "report-assets" / "personal-report"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
CODE_FONT = "Consolas"
ACCENT = "0F6B63"
BLUE = "1F4E78"
LIGHT_BLUE = "DDEBF7"
LIGHT_GRAY = "F4F6F8"
LIGHT_ORANGE = "FFF4D6"
GRID = "AAB7C4"

FLOW_TITLES = {
    1: "桥梁档案综合查询流程",
    2: "检查趋势数据聚合与展示流程",
    3: "检查到期提醒自动生成任务流程",
    4: "PDF 报告生成流程",
    5: "用户管理与安全保护流程",
    6: "操作日志记录与 IP 清洗流程",
    7: "版本控制与数据库备份流程",
}

FLOW_FILES = {index: ASSET_DIR / f"flow-{index:02d}.png" for index in FLOW_TITLES}

SCREENSHOT_FILES = [
    "UI-01-桥梁地图定位查询.png",
    "UI-02-检查趋势汇总.png",
    "UI-03-PDF生成工具栏.png",
    "UI-04-桥梁档案E表F表.png",
    "UI-05-用户管理列表.png",
    "UI-06-用户新增编辑抽屉.png",
    "UI-07-操作日志列表.png",
    "UI-08-版本控制与数据库备份.png",
    "UI-09-创建数据库备份弹窗.png",
    "UI-10-角色管理列表.png",
]


def font_path() -> Path:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    return next((item for item in candidates if item.exists()), candidates[-1])


CHART_FONT = FontProperties(fname=str(font_path()))


def node(node_id: str, x: float, y: float, width: float, height: float, text: str,
         kind: str = "process") -> dict:
    return {
        "id": node_id,
        "x": x,
        "y": y,
        "w": width,
        "h": height,
        "text": text,
        "kind": kind,
    }


def wrap_chart_text(value: str, max_units: int) -> str:
    lines = []
    for source_line in value.split("\n"):
        current = ""
        units = 0.0
        for char in source_line:
            char_units = 0.55 if ord(char) < 128 else 1.0
            if current and units + char_units > max_units:
                lines.append(current)
                current = char
                units = char_units
            else:
                current += char
                units += char_units
        if current:
            lines.append(current)
    return "\n".join(lines)


def anchor(source: dict, target: dict) -> tuple[float, float]:
    dx = target["x"] - source["x"]
    dy = target["y"] - source["y"]
    if abs(dx) < 1e-6 and abs(dy) < 1e-6:
        return source["x"], source["y"]
    half_w = source["w"] / 2
    half_h = source["h"] / 2
    if source["kind"] in {"start", "end"}:
        scale = 1 / math.sqrt((dx / half_w) ** 2 + (dy / half_h) ** 2)
    else:
        scales = []
        if abs(dx) > 1e-6:
            scales.append(half_w / abs(dx))
        if abs(dy) > 1e-6:
            scales.append(half_h / abs(dy))
        scale = min(scales)
    return source["x"] + dx * scale, source["y"] + dy * scale


def draw_node(axis, item: dict) -> None:
    x, y, width, height = item["x"], item["y"], item["w"], item["h"]
    kind = item["kind"]
    fills = {
        "start": "#DFF2E1",
        "end": "#DFF2E1",
        "process": "#EAF2F8",
        "decision": "#FFF0C7",
        "output": "#E2F1ED",
        "warning": "#FCE4DF",
    }
    edge = {
        "start": "#3A7D44",
        "end": "#3A7D44",
        "process": "#315B7D",
        "decision": "#A66A00",
        "output": "#0F766E",
        "warning": "#B44634",
    }[kind]
    if kind in {"start", "end"}:
        patch = Ellipse((x, y), width, height, facecolor=fills[kind], edgecolor=edge, linewidth=1.6)
    elif kind == "decision":
        patch = Polygon(
            [(x, y + height / 2), (x + width / 2, y), (x, y - height / 2), (x - width / 2, y)],
            closed=True,
            facecolor=fills[kind],
            edgecolor=edge,
            linewidth=1.6,
        )
    else:
        patch = FancyBboxPatch(
            (x - width / 2, y - height / 2),
            width,
            height,
            boxstyle="round,pad=0.25,rounding_size=1.2",
            facecolor=fills[kind],
            edgecolor=edge,
            linewidth=1.5,
        )
    axis.add_patch(patch)
    max_units = max(7, int(width * (0.58 if kind == "decision" else 0.72)))
    axis.text(
        x,
        y,
        wrap_chart_text(item["text"], max_units),
        ha="center",
        va="center",
        fontproperties=CHART_FONT,
        fontsize=9.2,
        color="#172033",
        linespacing=1.25,
        zorder=3,
    )


def draw_chart(index: int, nodes: list[dict], edges: list[tuple]) -> None:
    path = FLOW_FILES[index]
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, axis = plt.subplots(figsize=(10, 7), dpi=200)
    fig.patch.set_facecolor("white")
    axis.set_xlim(0, 100)
    axis.set_ylim(0, 100)
    axis.axis("off")
    axis.text(
        50,
        97,
        FLOW_TITLES[index],
        ha="center",
        va="center",
        fontproperties=CHART_FONT,
        fontsize=17,
        fontweight="bold",
        color="#17324D",
    )
    lookup = {item["id"]: item for item in nodes}
    for edge in edges:
        source_id, target_id = edge[0], edge[1]
        label = edge[2] if len(edge) > 2 else ""
        curve = edge[3] if len(edge) > 3 else 0.0
        source = lookup[source_id]
        target = lookup[target_id]
        start = anchor(source, target)
        finish = anchor(target, source)
        arrow = FancyArrowPatch(
            start,
            finish,
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.35,
            color="#5C6B78",
            connectionstyle=f"arc3,rad={curve}",
            zorder=1,
        )
        axis.add_patch(arrow)
        if label:
            mid_x = (start[0] + finish[0]) / 2
            mid_y = (start[1] + finish[1]) / 2
            axis.text(
                mid_x,
                mid_y + 1.2,
                label,
                ha="center",
                va="center",
                fontproperties=CHART_FONT,
                fontsize=8.5,
                color="#43515D",
                bbox={"facecolor": "white", "edgecolor": "none", "pad": 1.3},
                zorder=2,
            )
    for item in nodes:
        draw_node(axis, item)
    axis.text(
        99,
        1,
        "公路桥梁初始检查信息系统",
        ha="right",
        va="bottom",
        fontproperties=CHART_FONT,
        fontsize=7.5,
        color="#7A8793",
    )
    fig.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
    plt.close(fig)


def generate_flowcharts() -> None:
    charts: dict[int, tuple[list[dict], list[tuple]]] = {}

    charts[1] = (
        [
            node("start", 50, 90, 24, 7, "进入查询统计模块", "start"),
            node("mode", 50, 77, 20, 10, "选择查询方式", "decision"),
            node("map_in", 21, 63, 27, 8, "输入地图关键词\n桥号/名称/地址"),
            node("list_in", 79, 63, 29, 8, "设置组合条件\n路线/桥型/年份等"),
            node("map_api", 21, 49, 30, 8, "GET /api/bridge-profiles/map-points"),
            node("list_api", 79, 49, 28, 8, "GET /api/bridges\npage + filters"),
            node("map_service", 21, 34, 31, 9, "查询有效坐标\n统计桥墩并转换 GCJ-02"),
            node("list_service", 79, 34, 31, 9, "动态拼接参数化 SQL\nCOUNT + LIMIT/OFFSET"),
            node("map_out", 21, 19, 31, 8, "渲染 Marker、侧栏\n搜索后自动缩放", "output"),
            node("list_out", 79, 19, 31, 8, "渲染表格与分页\n保留筛选状态", "output"),
            node("profile", 50, 7, 30, 7, "点击结果进入桥梁档案", "end"),
        ],
        [
            ("start", "mode"),
            ("mode", "map_in", "地图查询"),
            ("mode", "list_in", "列表查询"),
            ("map_in", "map_api"),
            ("map_api", "map_service"),
            ("map_service", "map_out"),
            ("list_in", "list_api"),
            ("list_api", "list_service"),
            ("list_service", "list_out"),
            ("map_out", "profile"),
            ("list_out", "profile"),
        ],
    )

    charts[2] = (
        [
            node("start", 50, 90, 27, 7, "用户点击 Σ 汇总卡片", "start"),
            node("api", 50, 78, 32, 8, "GET /api/bridge-profiles/{bridgeCode}"),
            node("initial", 25, 63, 32, 9, "查询全部初始检查\n按有效标志和日期排序"),
            node("periodic", 75, 63, 32, 9, "查询全部定期检查\n关联技术状况等级"),
            node("initial_map", 25, 48, 31, 8, "构造 initialHistory\n日期/版本/有效状态"),
            node("batch", 75, 48, 33, 9, "一次 IN 查询全部部件记录\n按检查编号分组"),
            node("aggregate", 50, 33, 34, 9, "聚合检查次数、等级\n过滤“未见明显缺损”并计数"),
            node("json", 50, 19, 33, 8, "返回 inspectionSummary JSON", "output"),
            node("render", 50, 7, 37, 7, "渲染双柱状图与两张历史表", "end"),
        ],
        [
            ("start", "api"),
            ("api", "initial"),
            ("api", "periodic"),
            ("initial", "initial_map"),
            ("periodic", "batch"),
            ("initial_map", "aggregate"),
            ("batch", "aggregate"),
            ("aggregate", "json"),
            ("json", "render"),
        ],
    )

    charts[3] = (
        [
            node("start", 50, 91, 30, 7, "每天 00:05 定时触发", "start"),
            node("range", 50, 79, 33, 8, "计算 today 与 today+1个月"),
            node("query", 50, 65, 37, 9, "查询最新定期检查到期记录\n无定检时回退有效初始检查"),
            node("candidate", 50, 51, 22, 10, "存在候选桥梁？", "decision"),
            node("none", 82, 37, 25, 7, "返回 created=0", "end"),
            node("create", 43, 37, 37, 9, "生成 JC 任务编号\n计划期=到期前1个月至到期日"),
            node("insert", 43, 23, 36, 8, "写入待分配任务及状态历史"),
            node("more", 43, 11, 22, 9, "还有候选？", "decision"),
            node("result", 75, 7, 30, 7, "返回创建数量与截止日期", "end"),
        ],
        [
            ("start", "range"),
            ("range", "query"),
            ("query", "candidate"),
            ("candidate", "none", "否"),
            ("candidate", "create", "是"),
            ("create", "insert"),
            ("insert", "more"),
            ("more", "create", "", 0.42),
            ("more", "result", "否"),
        ],
    )

    charts[4] = (
        [
            node("start", 50, 91, 30, 7, "点击当前卡片的生成 PDF 按钮", "start"),
            node("auth", 50, 80, 34, 7, "JWT 鉴权：admin / engineer"),
            node("type", 50, 68, 22, 10, "报告类型？", "decision"),
            node("a", 12, 52, 20, 9, "A 表\n桥梁基本状况"),
            node("b", 37, 52, 20, 9, "B 表\n有效初始检查"),
            node("c", 63, 52, 20, 9, "C 表\n指定定期检查"),
            node("s", 88, 52, 20, 9, "Σ 汇总\n历次检查趋势"),
            node("validate", 50, 37, 36, 9, "查询关联数据并校验记录存在性"),
            node("html", 50, 25, 35, 8, "组装规范化 HTML 与中文字体"),
            node("pdf", 50, 14, 36, 8, "OpenHTMLToPDF 渲染并写入 reports/", "output"),
            node("done", 50, 4, 39, 7, "登记 tb_report，返回下载并记录日志", "end"),
        ],
        [
            ("start", "auth"),
            ("auth", "type"),
            ("type", "a", "A"),
            ("type", "b", "B"),
            ("type", "c", "C"),
            ("type", "s", "Σ"),
            ("a", "validate"),
            ("b", "validate"),
            ("c", "validate"),
            ("s", "validate"),
            ("validate", "html"),
            ("html", "pdf"),
            ("pdf", "done"),
        ],
    )

    charts[5] = (
        [
            node("create_start", 27, 91, 27, 7, "管理员新增用户", "start"),
            node("form", 27, 79, 31, 8, "校验姓名、账号、角色与状态"),
            node("pwd", 27, 66, 22, 10, "填写初始密码？", "decision"),
            node("given", 13, 52, 22, 8, "使用管理员输入值"),
            node("random", 41, 52, 23, 8, "生成 12 位随机值"),
            node("hash", 27, 39, 28, 8, "BCrypt 加密密码"),
            node("insert", 27, 27, 31, 8, "写入 tb_user\n默认 force_pwd_change=1"),
            node("login", 27, 14, 33, 8, "登录返回 forcePwdChange\n后端改密接口可用", "output"),
            node("ui_gap", 27, 4, 36, 6, "前端改密路由/视图待补齐", "warning"),
            node("delete_start", 76, 91, 27, 7, "管理员请求删除用户", "start"),
            node("self", 76, 77, 21, 10, "删除自己？", "decision"),
            node("reject_self", 57, 63, 24, 7, "拒绝：不能删除自己", "warning"),
            node("last", 82, 62, 25, 10, "最后管理员？", "decision"),
            node("reject_last", 67, 47, 27, 8, "拒绝：保留最后管理员", "warning"),
            node("soft", 91, 47, 17, 8, "软删除\nis_deleted=1", "output"),
            node("delete_end", 82, 33, 27, 7, "返回操作结果并记日志", "end"),
        ],
        [
            ("create_start", "form"),
            ("form", "pwd"),
            ("pwd", "given", "是"),
            ("pwd", "random", "否"),
            ("given", "hash"),
            ("random", "hash"),
            ("hash", "insert"),
            ("insert", "login"),
            ("login", "ui_gap"),
            ("delete_start", "self"),
            ("self", "reject_self", "是"),
            ("self", "last", "否"),
            ("last", "reject_last", "是"),
            ("last", "soft", "否"),
            ("reject_self", "delete_end"),
            ("reject_last", "delete_end"),
            ("soft", "delete_end"),
        ],
    )

    charts[6] = (
        [
            node("start", 50, 91, 32, 7, "用户执行关键业务操作", "start"),
            node("call", 50, 79, 38, 8, "Controller 调用 OperationLogService.log"),
            node("identity", 50, 66, 35, 8, "从 SecurityContext 读取用户标识"),
            node("xff", 50, 52, 24, 10, "存在 X-Forwarded-For？", "decision"),
            node("proxy", 25, 38, 31, 8, "取最左侧代理 IP"),
            node("remote", 75, 38, 31, 8, "读取 request.getRemoteAddr"),
            node("clean", 50, 25, 39, 9, "清洗 ::ffff: 前缀\nIPv6 回环转换为 127.0.0.1"),
            node("insert", 50, 13, 39, 8, "写入 tb_operation_log\n用户/IP/模块/对象/结果", "output"),
            node("end", 50, 3, 34, 6, "日志页只读分页查询", "end"),
        ],
        [
            ("start", "call"),
            ("call", "identity"),
            ("identity", "xff"),
            ("xff", "proxy", "是"),
            ("xff", "remote", "否"),
            ("proxy", "clean"),
            ("remote", "clean"),
            ("clean", "insert"),
            ("insert", "end"),
        ],
    )

    charts[7] = (
        [
            node("start", 50, 91, 30, 7, "管理员进入版本与备份页", "start"),
            node("action", 50, 78, 22, 10, "选择操作", "decision"),
            node("check", 24, 64, 31, 8, "检查正式版本\ngit fetch --prune --tags"),
            node("backup", 76, 64, 30, 8, "创建数据库备份"),
            node("compare", 24, 50, 32, 9, "读取 vX.Y 标签并比较\n已安装提交与最新提交"),
            node("tables", 76, 50, 32, 9, "SHOW FULL TABLES\n逐表读取结构和数据"),
            node("update", 24, 36, 22, 10, "执行更新/回溯？", "decision"),
            node("dump", 76, 36, 34, 9, "生成 DROP/CREATE/INSERT\nUTF-8 SQL 文件"),
            node("summary", 10, 21, 19, 8, "仅返回版本摘要", "output"),
            node("safety", 37, 21, 32, 9, "检查工作区清洁\n安全分支 + 数据库备份 + reset"),
            node("hash", 76, 21, 31, 8, "计算 SHA-256 并登记元数据"),
            node("end", 50, 6, 38, 7, "返回结果、刷新列表并记录审计日志", "end"),
        ],
        [
            ("start", "action"),
            ("action", "check", "版本"),
            ("action", "backup", "备份"),
            ("check", "compare"),
            ("compare", "update"),
            ("update", "summary", "否"),
            ("update", "safety", "是"),
            ("backup", "tables"),
            ("tables", "dump"),
            ("dump", "hash"),
            ("summary", "end"),
            ("safety", "end"),
            ("hash", "end"),
        ],
    )

    for index, (nodes, edges) in charts.items():
        draw_chart(index, nodes, edges)


def set_east_asia(run_or_style, name: str) -> None:
    element = run_or_style._element
    if element.tag == qn("w:r"):
        properties = element.get_or_add_rPr()
    else:
        properties = element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)


def set_run_font(run, name: str = BODY_FONT, size: float = 12, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    set_east_asia(run, name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_or_create(document: Document, name: str, base: str = "Normal"):
    try:
        return document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles[base]
        return style


def add_outline_level(style, level: int) -> None:
    paragraph_properties = style._element.get_or_add_pPr()
    existing = paragraph_properties.find(qn("w:outlineLvl"))
    if existing is not None:
        paragraph_properties.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    paragraph_properties.append(outline)


def configure_styles(document: Document) -> None:
    body = style_or_create(document, "Report Body")
    body.font.name = BODY_FONT
    body.font.size = Pt(12)
    set_east_asia(body, BODY_FONT)
    body.paragraph_format.first_line_indent = Cm(0.74)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.widow_control = True

    bullet = style_or_create(document, "Report Bullet", "Report Body")
    bullet.paragraph_format.first_line_indent = Cm(0)
    bullet.paragraph_format.left_indent = Cm(0.74)
    bullet.paragraph_format.hanging_indent = Cm(0.37)
    bullet.paragraph_format.line_spacing = 1.35

    heading_specs = [
        ("Report Heading 1", 0, 16, WD_ALIGN_PARAGRAPH.CENTER, 18, 12),
        ("Report Heading 2", 1, 14, WD_ALIGN_PARAGRAPH.LEFT, 14, 6),
        ("Report Heading 3", 2, 12, WD_ALIGN_PARAGRAPH.LEFT, 10, 4),
        ("Report Heading 4", 3, 11, WD_ALIGN_PARAGRAPH.LEFT, 8, 3),
    ]
    for name, level, size, alignment, before, after in heading_specs:
        style = style_or_create(document, name)
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("172033")
        set_east_asia(style, HEADING_FONT)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.2
        add_outline_level(style, level)

    ui_heading = style_or_create(document, "Report UI Heading")
    ui_heading.font.name = HEADING_FONT
    ui_heading.font.size = Pt(12)
    ui_heading.font.bold = True
    ui_heading.font.color.rgb = RGBColor.from_string(ACCENT)
    set_east_asia(ui_heading, HEADING_FONT)
    ui_heading.paragraph_format.space_before = Pt(10)
    ui_heading.paragraph_format.space_after = Pt(4)
    ui_heading.paragraph_format.keep_with_next = True

    caption = style_or_create(document, "Report Caption")
    caption.font.name = BODY_FONT
    caption.font.size = Pt(10.5)
    set_east_asia(caption, BODY_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    code = style_or_create(document, "Report Code")
    code.font.name = CODE_FONT
    code.font.size = Pt(8)
    set_east_asia(code, BODY_FONT)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_after = Pt(0)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: int = 6) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        item = borders.find(tag)
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), str(size))
        item.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node_value = margins.find(qn(f"w:{key}"))
        if node_value is None:
            node_value = OxmlElement(f"w:{key}")
            margins.append(node_value)
        node_value.set(qn("w:w"), str(value))
        node_value.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def clean_text(value: str) -> str:
    return (
        html.unescape(value)
        .replace("\\|", "|")
        .replace("📷", "")
        .replace("✅", "")
        .replace("&emsp;", "  ")
        .strip()
    )


def add_inline(paragraph, value: str, size: float = 12, inherited_bold: bool = False) -> None:
    value = clean_text(value)
    index = 0
    plain_start = 0

    def flush(end: int) -> None:
        nonlocal plain_start
        if end <= plain_start:
            return
        run = paragraph.add_run(value[plain_start:end])
        set_run_font(run, BODY_FONT, size, bold=inherited_bold)

    while index < len(value):
        if value.startswith("**", index):
            close = value.find("**", index + 2)
            if close >= 0:
                flush(index)
                add_inline(paragraph, value[index + 2:close], size=size, inherited_bold=True)
                index = close + 2
                plain_start = index
                continue
        if value[index] == "`":
            close = value.find("`", index + 1)
            if close >= 0:
                flush(index)
                run = paragraph.add_run(value[index + 1:close])
                set_run_font(run, CODE_FONT, max(8.5, size - 1), bold=inherited_bold, color="7A3E00")
                index = close + 1
                plain_start = index
                continue
        break_match = re.match(r"<br\s*/?>", value[index:], flags=re.IGNORECASE)
        if break_match:
            flush(index)
            paragraph.add_run().add_break()
            index += len(break_match.group(0))
            plain_start = index
            continue
        link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", value[index:])
        if link_match:
            flush(index)
            display, url = link_match.groups()
            run = paragraph.add_run(f"{display}（{url}）")
            set_run_font(run, BODY_FONT, size, bold=inherited_bold, color=BLUE)
            run.underline = True
            index += len(link_match.group(0))
            plain_start = index
            continue
        index += 1
    flush(len(value))


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in table_lines]
    if len(rows) > 1 and is_table_separator(rows[1]):
        rows.pop(1)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    rows = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    widths = {
        2: [3.2, 12.2],
        3: [3.0, 5.2, 7.2],
        4: [2.2, 3.0, 4.3, 5.9],
        5: [2.1, 2.7, 3.0, 4.4, 3.2],
        6: [1.7, 2.4, 2.5, 3.3, 2.4, 3.1],
    }.get(column_count)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            repeat_table_header(row)
        else:
            prevent_row_split(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            if widths:
                cell.width = Cm(widths[column_index])
            paragraph = cell.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value, size=8.5, inherited_bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
    document.add_paragraph(style="Report Body").paragraph_format.space_after = Pt(2)


def add_code_block(document: Document, lines: list[str], language: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_border(cell, "C6CDD4", 5)
    set_cell_margins(cell, 90, 120, 90, 120)
    first = cell.paragraphs[0]
    clear_paragraph(first)
    for line_index, line in enumerate(lines or [""]):
        paragraph = first if line_index == 0 else cell.add_paragraph()
        paragraph.style = document.styles["Report Code"]
        run = paragraph.add_run(line.rstrip())
        set_run_font(run, CODE_FONT, 8)
    document.add_paragraph(style="Report Body").paragraph_format.space_after = Pt(2)


def add_callout(document: Document, value: str, fill: str = "EAF4F2") -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, ACCENT, 7)
    set_cell_margins(cell, 100, 150, 100, 150)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.25
    add_inline(paragraph, value, size=10.5)
    document.add_paragraph(style="Report Body").paragraph_format.space_after = Pt(2)


def screenshot_label(description: str) -> str:
    text = clean_text(re.sub(r"\*\*", "", description))
    text = re.sub(r"^截图说明[：:]?", "", text).strip(" ：:")
    return text


def add_screenshot_placeholder(document: Document, description: str, index: int,
                               interface_title: str) -> None:
    figure_number = 7 + index
    file_name = SCREENSHOT_FILES[index - 1] if index <= len(SCREENSHOT_FILES) else f"UI-{index:02d}.png"
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Cm(2.8)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    shade_cell(cell, LIGHT_ORANGE)
    set_cell_border(cell, "D39B2A", 10)
    set_cell_margins(cell, 160, 190, 160, 190)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    title = paragraph.add_run(f"【此处插入图2-{figure_number}：{interface_title or '系统界面截图'}】")
    set_run_font(title, HEADING_FONT, 11, bold=True, color="8A5A00")
    detail = cell.add_paragraph()
    detail.paragraph_format.first_line_indent = Cm(0)
    detail.paragraph_format.line_spacing = 1.2
    add_inline(detail, f"截图内容：{screenshot_label(description)}", size=10)
    file_paragraph = cell.add_paragraph()
    file_paragraph.paragraph_format.first_line_indent = Cm(0)
    file_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = file_paragraph.add_run(f"建议文件名：{file_name}；建议分辨率：1280×720 或更高；请隐藏账号、密钥等敏感信息。")
    set_run_font(run, BODY_FONT, 9.5, color="6E4B00", italic=True)
    caption = document.add_paragraph(style="Report Caption")
    run = caption.add_run(f"图2-{figure_number} {interface_title or '系统界面'}（待替换为系统真实运行截图）")
    set_run_font(run, BODY_FONT, 10.5)


def add_flow_figure(document: Document, index: int) -> None:
    path = FLOW_FILES[index]
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Cm(15.6))
    shape._inline.docPr.set("descr", FLOW_TITLES[index])
    caption = document.add_paragraph(style="Report Caption")
    caption_run = caption.add_run(f"图2-{index} {FLOW_TITLES[index]}")
    set_run_font(caption_run, BODY_FONT, 10.5)


def add_heading(document: Document, text: str, level: int, first_chapter: bool) -> bool:
    if level == 1 and not first_chapter:
        document.add_page_break()
    paragraph = document.add_paragraph(style=f"Report Heading {level}")
    run = paragraph.add_run(clean_text(text))
    set_run_font(run, HEADING_FONT, {1: 16, 2: 14, 3: 12, 4: 11}[level], bold=True)
    return first_chapter or level == 1


def is_special_line(line: str) -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("```")
        or stripped.startswith(">")
        or stripped.startswith("|")
        or bool(re.match(r"^[-*]\s+", stripped))
        or stripped == "---"
    )


def render_markdown(document: Document, source: str) -> tuple[int, int]:
    lines = source.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 第一章"))
    lines = lines[start:]
    index = 0
    first_chapter = True
    pending_flow: int | None = None
    screenshot_count = 0
    interface_title = ""

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped or stripped == "---":
            index += 1
            continue

        heading_match = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            if len(hashes) == 5 and title.startswith("流程图"):
                number_match = re.search(r"流程图\s*(\d+)", title)
                pending_flow = int(number_match.group(1)) if number_match else None
                index += 1
                continue
            level = max(1, len(hashes) - 1)
            first_chapter = add_heading(document, title, level, first_chapter)
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            block = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            if pending_flow is not None:
                add_flow_figure(document, pending_flow)
                pending_flow = None
            else:
                add_code_block(document, block, language)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(re.sub(r"^>\s?", "", lines[index].strip()))
                index += 1
            quote = " ".join(quote_lines)
            if "截图说明" in quote:
                screenshot_count += 1
                add_screenshot_placeholder(document, quote, screenshot_count, interface_title)
            else:
                add_callout(document, quote)
            continue

        if re.match(r"^[-*]\s+", stripped):
            while index < len(lines) and re.match(r"^[-*]\s+", lines[index].strip()):
                content = re.sub(r"^[-*]\s+", "", lines[index].strip())
                paragraph = document.add_paragraph(style="Report Bullet")
                bullet_run = paragraph.add_run("• ")
                set_run_font(bullet_run, BODY_FONT, 12)
                add_inline(paragraph, content, size=12)
                index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not is_special_line(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        text = " ".join(paragraph_lines)
        plain = clean_text(re.sub(r"\*\*", "", text))
        if plain.startswith("【界面"):
            interface_title = re.sub(r"^【界面\s*\d+】", "", plain).strip()
            paragraph = document.add_paragraph(style="Report UI Heading")
            run = paragraph.add_run(plain)
            set_run_font(run, HEADING_FONT, 12, bold=True, color=ACCENT)
        else:
            paragraph = document.add_paragraph(style="Report Body")
            if text.startswith("**") and text.endswith("**"):
                paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, text, size=12)

    return screenshot_count, 7


def truncate_template(document: Document) -> None:
    body = document.element.body
    deleting = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        text = "".join(child.itertext()).replace(" ", "")
        if child.tag == qn("w:p") and "目录" in text:
            deleting = True
        if deleting:
            body.remove(child)


def set_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, BODY_FONT, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instruction)
    field_run.append(separate)
    field_run.append(result)
    field_run.append(end)
    paragraph._p.append(field_run)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, BODY_FONT, 9)

    section_properties = section._sectPr
    page_number_type = section_properties.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section_properties.append(page_number_type)
    page_number_type.set(qn("w:start"), "1")


def add_toc(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("目  录")
    set_run_font(run, HEADING_FONT, 18, bold=True)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后将自动更新目录；如页码未刷新，请按 Ctrl+A、F9。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instruction)
    field_run.append(separate)
    field_run.append(placeholder)
    field_run.append(end)
    paragraph._p.append(field_run)
    document.add_page_break()


def enable_field_update(document: Document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def fill_cover(document: Document) -> None:
    values = [
        ("题    目", "公路桥梁初始检查信息系统设计与实现"),
        ("课程名称", "程序设计综合实践II"),
        ("专业班级", "计算机科学与技术24级曙光1班"),
        ("团队名称", "雪碧&巧乐兹"),
        ("学    号", "632407070128"),
        ("姓    名", "肖皓轩"),
        ("指导教师", "王家伟"),
    ]
    cover = document.tables[0]
    for row, (label, value) in zip(cover.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    for row in cover.rows:
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, BODY_FONT, 14 if cell_index == 1 else 12, bold=cell_index == 0)


def build_document() -> tuple[int, int]:
    generate_flowcharts()
    document = Document(str(TARGET_DOCX))
    fill_cover(document)
    truncate_template(document)
    configure_styles(document)

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    set_page_number(section)
    add_toc(document)

    source = SOURCE_MD.read_text(encoding="utf-8")
    screenshot_count, flow_count = render_markdown(document, source)
    enable_field_update(document)

    document.core_properties.title = "公路桥梁初始检查信息系统设计与实现——成员设计报告"
    document.core_properties.subject = "M05 查询统计与决策支持；M06 系统管理"
    document.core_properties.author = "肖皓轩"
    document.core_properties.last_modified_by = "肖皓轩"
    document.core_properties.keywords = "桥梁初始检查, 查询统计, 系统管理, 课程设计"

    temp_path = TARGET_DOCX.with_name(TARGET_DOCX.stem + ".generated.docx")
    document.save(str(temp_path))
    os.replace(temp_path, TARGET_DOCX)
    return screenshot_count, flow_count


if __name__ == "__main__":
    screenshots, flows = build_document()
    print(f"Generated: {TARGET_DOCX}")
    print(f"Flowcharts: {flows}; screenshot placeholders: {screenshots}")
