# -*- coding: utf-8 -*-
"""Generate the six-subsystem overview design from the supplied Word template."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "bridge-inspection-system"
TEMPLATE = ROOT / "4. 概要设计示例.docx"
OUTPUT = ROOT / "公路桥梁检查信息系统概要设计.docx"
FUNCTION_DIAGRAM = ROOT / "系统总体功能设计图.png"
MODEL_ASSETS = PROJECT / "docs" / "overview-assets" / "powerdesigner"
CDM_DIAGRAM = MODEL_ASSETS / "PowerDesigner_CDM_桥梁核心结构关系.png"
LDM_DIAGRAM = MODEL_ASSETS / "PowerDesigner_LDM_桥梁核心结构关系.png"
PDM_DIAGRAM = MODEL_ASSETS / "PowerDesigner_PDM_桥梁核心结构关系.png"


FUNCTION_MODULES = [
    (
        "3.1.1 系统管理子系统",
        [
            (
                "角色设置及权限配置",
                "完成角色的新增、修改、查询和权限配置。系统以角色为授权单位，将菜单访问、数据查看、数据维护、审核和系统管理等权限分配给角色，并由前端路由与后端接口共同校验。系统预置系统管理员、桥梁工程师、检查人员、审核人员和查询人员五类角色。",
                "角色数据包括角色编号、角色代码、角色名称、角色说明、权限集合和创建时间。角色代码与角色名称均须唯一；一个角色可关联多个用户。权限集合采用结构化数据保存，角色停用前须检查是否仍有关联用户，权限调整写入操作日志。",
            ),
            (
                "用户管理",
                "完成用户新增、编辑、查询、启用、停用、角色分配、密码重置和登录状态管理。公开注册的账号默认授予查询人员角色，管理员可根据职责调整角色；不同角色登录后进入相应工作台。",
                "用户数据包括用户编号、用户姓名、登录账号、密码密文、角色编号、单位部门、联系电话、电子邮箱、用户状态、是否强制修改密码、创建时间和最后登录时间。登录账号全局唯一，密码采用BCrypt散列保存，停用用户不得登录系统。",
            ),
            (
                "系统日志查看",
                "记录用户登录、退出、业务数据新增修改删除、任务流转、报告生成、文件上传和系统配置等关键操作，并支持按用户、模块、操作类型、结果和时间范围查询，为问题追踪和责任审计提供依据。",
                "日志数据包括日志编号、用户编号与姓名、业务模块、操作类型、操作说明、请求方法、请求路径、客户端地址、执行结果、异常摘要和操作时间。日志原则上只追加不修改，按操作时间和用户编号建立索引，普通业务用户无权删除日志。",
            ),
            (
                "数据备份与恢复",
                "提供数据库备份任务的创建、执行结果登记、备份文件校验、下载和恢复。恢复前应校验备份文件完整性和数据库版本，并要求管理员二次确认；恢复操作必须保留完整审计记录。",
                "备份数据包括备份编号、文件名称、存储路径、文件大小、备份类型、备份状态、执行人、开始时间、完成时间和结果说明。备份文件存放在受控目录，数据库仅保存元数据；恢复失败不得覆盖当前有效数据。",
            ),
        ],
    ),
    (
        "3.1.2 基础信息管理子系统",
        [
            (
                "路线信息管理",
                "完成公路路线的新增、修改、删除和查询。路线是桥梁归属、检查任务组织和统计分析的基础维度，一条路线可以关联多座桥梁。",
                "路线数据包括路线编号、路线名称、路线等级、创建时间和更新时间。路线编号作为业务主键且全局唯一；已被桥梁引用的路线不得直接物理删除，可通过状态字段停用。",
            ),
            (
                "基础字典维护",
                "统一维护桥梁类型、结构部位、部件、检查类别、缺损程度、技术状况等级、档案资料项和初始检查项目等基础字典，并配置桥型与部件、桥型与初检项目的适用矩阵。",
                "字典数据主要包括编码、名称、说明、排序号和启用状态。系统维护6类桥型、部位与部件关系、47个初检项目及其适用条件；同类编码不得重复。部件必须归属一个部位，桥型配置按桥型、部位、部件或检查项目建立组合唯一约束。",
            ),
            (
                "桥梁基础信息管理",
                "完成桥梁行政识别信息、桥位信息、结构类型、技术指标、建设与管养单位以及地图坐标的新增、修改、查询和状态管理。桥梁是档案、检查、病害、任务和报告的核心关联对象。",
                "桥梁数据包括桥梁编号、路线编号、桥梁类型编码、桥梁名称、行政区划、桥位桩号、详细地址、经纬度、管理机构、功能类型、被跨越道路、养护检查等级、设计荷载、桥梁全长、桥面总宽、车道与人行道宽度、净空、跨径组合、建成年份和各参建单位。桥梁编号全局唯一，删除采用逻辑删除。",
            ),
        ],
    ),
    (
        "3.1.3 桥梁档案管理子系统",
        [
            (
                "桥梁基础状况卡片管理",
                "按照一桥一档原则管理桥梁基本状况卡片（A表），汇总桥梁基础信息、结构部件、档案资料、检测评定历史和照片附件。结构信息按D-1至D-N分区展示，其中D-1固定展示桥墩、桥桩与基础实例，其余部件按实际桥型顺序形成后续分区。",
                "A表由桥梁主数据、桥梁具体部件、档案资料记录、检测评定历史和附件共同组成。每个实际桥墩、桥桩、基础、支座、主梁或索构件均保存为一条具体部件记录，数据包括部位、部件、序号、位置、材料、尺寸、数量、索力或内力、高程或变位、状态和备注。档案资料按资料项记录完整程度并关联电子文件。",
            ),
            (
                "桥梁档案查询",
                "支持按桥梁编号、桥梁名称、路线、桥型、行政区划、养护等级和档案完整程度查询桥梁档案，并可从列表或重庆桥梁地图进入桥梁详情，查看A表及其关联的初始检查、定期检查和报告。",
                "查询条件由桥梁主表和档案记录字段组成，结果包括桥梁摘要、档案完整程度、最近检查时间和技术状况等级。查询采用分页返回，查询人员只读访问；地图查询使用桥梁经纬度，附件通过受控文件接口预览或下载。",
            ),
            (
                "桥梁档案统计",
                "对桥梁档案按路线、桥型、行政区划、养护等级、建成年代和资料完整程度进行汇总，形成桥梁数量、结构类型分布和档案缺项统计，为档案补录和检查计划制定提供依据。",
                "统计数据由桥梁、桥梁类型、路线和档案记录聚合产生，至少包括分组维度、桥梁数量、占比、完整档案数、不完整档案数和缺失档案数。统计口径仅包含有效桥梁，并保留查询条件以便下钻到明细。",
            ),
        ],
    ),
    (
        "3.1.4 初始检测子系统",
        [
            (
                "初始检测记录录入",
                "依据JTG 5120-2021附录B录入桥梁初始检查。选择桥梁后，系统根据桥型初检项目矩阵生成适用检测项目，并允许按实际桥墩、桥桩等具体部件记录检测结果、病害说明和附件。初始记录经提交、审核后归档。",
                "初始检查主数据包括检查编号、桥梁编号、检查日期、检查机构、检查人员、桥梁工程师、气候与环境温度、主跨结构、最大跨径、跨径组合、施工方法、维修加固情况、病害描述及养护建议、状态和创建人。明细数据包括检测项目编码、检测结果、检测说明、具体部件编号、病害类型、位置、性质、范围、数量、程度和照片。每座桥梁只保留一份有效初始检查记录。",
            ),
            (
                "初始检测记录查询",
                "支持按检查编号、桥梁、路线、桥型、检查日期、检查机构、记录状态和病害条件查询初始检查记录，并查看公共检测项目、具体部件检测结果、病害、附件、任务和审核状态。",
                "查询主键为初始检查编号，桥梁编号作为唯一业务关联。列表返回检查日期、桥梁名称、检查机构、记录状态和病害摘要；详情聚合初始检查项目、具体部件检查、附件及报告。已归档记录按只读方式显示，修改须执行退回流程。",
            ),
            (
                "初始检测记录报表打印",
                "按照初始检查记录表版式汇总桥梁基础信息、检测项目、具体部件结果、病害、结论和签字信息，生成可预览、下载和打印的检查报告。报告与原始记录、任务和审核人保持可追溯关联。",
                "报告数据包括报告编号、关联任务、初始检查编号、报告类型、版本号、文件格式、文件路径、报告状态、生成时间、生成人、审核人和变更摘要。已审核版本不得覆盖，重新生成时增加版本号；打印内容使用归档时的桥梁关键字段快照。",
            ),
        ],
    ),
    (
        "3.1.5 定期检测子系统",
        [
            (
                "定期检测记录录入",
                "依据JTG 5120-2021附录C录入历次定期检查。系统根据桥梁类型和桥型部件矩阵自动生成应检部件，对桥面系、上部结构、下部结构和附属设施逐项记录缺损程度、评分和养护建议，并形成桥梁技术状况等级。",
                "定期检查主数据包括检查编号、桥梁编号、检查日期、上次检查日期、上次维修日期、气候环境、全桥清洁情况、预防及修复状况、下次检查日期、记录人、负责人、技术状况等级和状态。部件检查数据包括具体部件编号、部位、部件、缺损类型、位置、范围、缺损程度、最不利构件、评分、养护建议、是否需要特殊检查和附件。一座桥梁可以有多份定期检查记录。",
            ),
            (
                "定期检测记录查询",
                "支持按检查编号、桥梁、路线、桥型、检查日期范围、技术状况等级、记录状态、部位和缺损类型查询定期检查记录，查看每次检查的部件明细、病害照片、审核过程和报告。",
                "查询以定期检查编号为主键，通过桥梁编号关联桥梁档案，通过具体部件编号关联实际构件。列表返回检查日期、技术状况等级、负责人、状态和病害数量；详情按D-1至D-N或结构部位组织部件检查记录，并采用分页和索引保证查询性能。",
            ),
            (
                "定期检测记录对比分析",
                "对同一桥梁不同日期的定期检查进行横向与纵向比较，展示技术状况等级变化、同一具体部件评分变化、病害新增或扩展情况以及下次检查建议，辅助识别劣化趋势。",
                "对比数据至少包括桥梁编号、检查编号与日期、技术状况等级、部件编号、部件名称、评分、缺损类型、缺损程度、范围和数量。比较时以桥梁具体部件编号为稳定匹配键；缺少对应历史部件时标记为新增或已撤销，不直接参与差值计算。",
            ),
            (
                "定期检测记录报表打印",
                "按照对应桥型的定期检查记录表汇总桥梁信息、检查概况、部件检查与评分、病害照片、技术状况评定、养护建议和签字信息，生成可下载和打印的版本化报告。",
                "报告数据包括报告编号、定期检查编号、任务编号、桥型、报告类型、版本号、文件格式、文件路径、状态、生成与审核信息。报告生成时冻结本次检查及桥梁关键字段，图片按部件和病害编号排序，归档报告只能新增修订版本。",
            ),
        ],
    ),
    (
        "3.1.6 查询统计子系统",
        [
            (
                "综合查询",
                "对桥梁档案、初始检查、定期检查、病害、任务和报告进行统一检索。用户可按路线、桥梁、桥型、行政区划、日期、技术等级和记录状态组合筛选，并从查询结果进入桥梁详情或地图位置。",
                "查询条件通过受控字段映射生成，不直接拼接用户输入。结果统一返回分页信息、排序字段和业务摘要；不同角色按照权限隐藏不可访问的操作和敏感字段。桥梁地图查询读取经纬度，缺少有效坐标的数据不生成地图标注。",
            ),
            (
                "桥梁技术状况等级统计分析",
                "按路线、桥型、行政区划和时间范围统计1类至5类桥梁数量及占比，并结合历次定期检查展示等级变化，为重点桥梁筛查和检查计划安排提供数据支持。",
                "统计以已归档的最新定期检查为默认口径，数据包括统计日期、分组维度、技术状况等级、桥梁数量、占比和环比变化。用户切换历史时点时按检查日期截取当时最新有效记录，避免重复计数。",
            ),
            (
                "检测到期提醒",
                "根据桥梁养护检查等级、最近检查日期和记录中的下次检查日期识别即将到期、当天到期和已逾期桥梁，在角色工作台和查询页面形成提醒。",
                "提醒数据包括桥梁编号、桥梁名称、路线、养护检查等级、最近检查日期、计划下次检查日期、剩余天数和提醒状态。优先采用审核归档记录中的下次检查日期；无明确日期时按配置周期计算，同一桥梁同一检查类型只保留一条当前提醒。",
            ),
            (
                "数据导出",
                "将综合查询、桥梁清单、检查记录、病害明细和统计结果导出为Excel、CSV或PDF，并保留筛选条件、生成时间和导出人，满足离线汇报、交换和归档需求。",
                "导出任务保存导出类型、查询条件、字段清单、文件格式、文件路径、文件大小、生成状态、创建人和生成时间。导出字段受角色权限控制，大数据量采用后台生成；文件名使用随机标识，下载通过鉴权接口完成。",
            ),
        ],
    ),
]


def set_run_font(run, east_asia="宋体", size=None, bold=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_document(doc):
    doc.core_properties.title = "公路桥梁检查信息系统概要设计"
    doc.core_properties.subject = "《程序设计综合实践II》第3部分 系统设计"
    doc.core_properties.keywords = "公路桥梁, 概要设计, 功能设计, 数据模型"

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        style.font.bold = True

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.clear()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.add_run("《程序设计综合实践II》概要设计")
    set_run_font(header_run, "宋体", 9)
    paragraph_properties = header_paragraph._p.get_or_add_pPr()
    paragraph_borders = OxmlElement("w:pBdr")
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "4")
    bottom_border.set(qn("w:space"), "1")
    bottom_border.set(qn("w:color"), "auto")
    paragraph_borders.append(bottom_border)
    paragraph_properties.append(paragraph_borders)

    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_font(run, "黑体", bold=True)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.75)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12)
    return paragraph


def add_function_title(doc, index, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Cm(0.85)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{index}、{text}")
    set_run_font(run, "宋体", 12, True)
    return paragraph


def add_label(doc, text):
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12, True)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return paragraph


def add_picture(doc, path, width_cm):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    return paragraph


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, value, bold=False, center=False, size=9.5):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(value))
    set_run_font(run, "宋体", size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")
        borders.append(border)
    table._tbl.tblPr.append(borders)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    set_row_cant_split(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.width = Cm(widths[index])
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, center=True, size=10)
    for row_data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.width = Cm(widths[index])
            set_cell_text(cell, value)
    doc.add_paragraph()
    return table


def add_toc_field(doc):
    title = doc.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("目录")
    set_run_font(run, "黑体", 22, True)

    paragraph = doc.add_paragraph(style="Normal")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph.add_run()._r.append(begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    paragraph.add_run()._r.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    result = paragraph.add_run("第3部分 系统设计")
    set_run_font(result, "宋体", 12)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)
    doc.add_page_break()


def validate_inputs():
    required = [TEMPLATE, FUNCTION_DIAGRAM, CDM_DIAGRAM, LDM_DIAGRAM, PDM_DIAGRAM]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required files:\n" + "\n".join(missing))


def generate():
    validate_inputs()
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    configure_document(doc)
    add_toc_field(doc)

    add_heading(doc, "第3部分 系统设计", 1)
    add_body(
        doc,
        "以前期用户功能需求、系统总体功能设计图、JTG 5120-2021《公路桥涵养护规范》以及当前系统实现为基础，设计《公路桥梁检查信息系统》需要完成的功能和数据结构。系统以桥梁档案为核心，贯通初始检查、定期检查、病害、任务、审核、报告和查询统计，并通过统一的用户、角色与权限机制保障业务数据安全。",
    )

    add_heading(doc, "3.1 系统功能设计", 2)
    add_body(
        doc,
        "将用户功能需求转换为计算机软件系统需要实现的功能，整个《公路桥梁检查信息系统》由系统管理、基础信息管理、桥梁档案管理、初始检测、定期检测和查询统计六个子系统组成，系统总体功能架构如图3.1所示。",
    )
    add_picture(doc, FUNCTION_DIAGRAM, 14.5)
    add_caption(doc, "图3.1 系统总体功能架构")

    for subsystem_heading, functions in FUNCTION_MODULES:
        add_heading(doc, subsystem_heading, 3)
        for index, (name, description, data_description) in enumerate(functions, start=1):
            add_function_title(doc, index, name)
            add_label(doc, "功能描述")
            add_body(doc, description)
            add_label(doc, "操作数据描述")
            add_body(doc, data_description)

    doc.add_page_break()
    add_heading(doc, "3.2 数据模型设计", 2)
    add_body(
        doc,
        "数据模型以MySQL基线结构、桥梁模型示例图和PowerDesigner 16.5为依据，按概念数据模型、逻辑数据模型和物理数据模型三个层次描述。关系修正版包含32个业务对象、263个属性或字段、54条关系或外键，并按桥梁核心结构、基础字典、桥梁档案、初始检查、定期检查、任务报告和用户权限组织为6个子图。默认图严格展示部位、部件、桥梁类型部件配置、桥梁类型、桥梁、桥梁部件、桥梁具体部件、桥梁基本信息和桥梁具体部件检查记录九个核心实体。",
    )
    add_table(
        doc,
        ["模型层次", "正式模型文件", "主要规模"],
        [
            ["概念数据模型CDM", "公路桥梁初始检查信息系统概念模型_当前版.cdm", "32个实体、263个属性、54条关系、6个概念子图"],
            ["逻辑数据模型LDM", "公路桥梁初始检查信息系统逻辑模型_当前版.ldm", "32个实体、263个属性、54条关系、6个逻辑子图"],
            ["物理数据模型PDM", "公路桥梁初始检查信息系统物理模型_当前版.pdm", "32张表、263个字段、54条外键、69个索引、6个物理子图"],
        ],
        [3.2, 6.4, 4.9],
    )

    add_heading(doc, "3.2.1 概念数据模型", 3)
    add_body(
        doc,
        "概念模型将标准目录、桥型模板、桥梁结构实例和检查记录分为四个层次。部位包含标准部件；桥梁类型与标准部件通过桥梁类型部件配置形成模板；每座桥梁依据模板生成桥梁部件，再在桥梁部件下登记桥墩、支座、主梁等桥梁具体部件；检查结果必须关联到一个桥梁具体部件。桥梁基本信息独立保存版本，使一座桥梁可以保留多期基本信息，同时仅标记一条当前有效记录。",
    )
    add_table(
        doc,
        ["核心关系", "父实体可对应", "子实体必须对应", "设计含义"],
        [
            ["部位—部件", "1..N个部件", "1个部位", "每个部件只归属一个部位"],
            ["部件—桥梁类型部件配置", "1..N条配置", "1个部件", "配置引用一个标准部件"],
            ["桥梁类型—桥梁类型部件配置", "1..N条配置", "1个桥梁类型", "每种桥型至少定义一个部件配置"],
            ["桥梁类型—桥梁", "0..N座桥梁", "1个桥梁类型", "桥型可以暂未被桥梁采用"],
            ["桥梁类型部件配置—桥梁部件", "0..N个桥梁部件", "1条配置", "模板配置可实例化到多座桥梁"],
            ["桥梁—桥梁部件", "1..N个桥梁部件", "1座桥梁", "每座桥梁至少包含一个部件分组"],
            ["桥梁部件—桥梁具体部件", "1..N个具体部件", "1个桥梁部件", "同一部件分组可有多个实际实例"],
            ["桥梁—桥梁基本信息", "1..N条基本信息", "1座桥梁", "基本信息按版本保留历史"],
            ["桥梁具体部件—检查记录", "1..N条检查记录", "1个具体部件", "检查结果不可脱离具体部件"],
        ],
        [4.2, 3.0, 3.0, 4.3],
    )
    add_picture(doc, CDM_DIAGRAM, 14.2)
    add_caption(doc, "图3.2 PowerDesigner概念数据模型（桥梁核心结构关系）")

    add_heading(doc, "3.2.2 逻辑模型设计", 3)
    add_body(
        doc,
        "逻辑模型将九实体核心关系转换为满足第三范式的关系模式。tb_component通过part_code确定唯一部位；tb_bridge_type_component_config只关联桥梁类型和标准部件，不再重复关联部位；tb_bridge_component以桥梁编号和配置编号唯一标识某座桥梁的部件分组；tb_bridge_specific_component仅通过bridge_component_id归属桥梁部件，不再直接关联桥梁、桥型配置、部位和标准部件；tb_component_inspection仅通过bridge_specific_component_id确定被检查对象。",
    )
    add_table(
        doc,
        ["业务域", "主要逻辑表"],
        [
            ["系统管理", "tb_role、tb_user、tb_operation_log、tb_backup_record"],
            ["基础字典与矩阵", "tb_route、tb_bridge_type、tb_part、tb_component、tb_bridge_type_component_config、tb_initial_inspection_item_definition、tb_bridge_type_initial_item_config、tb_check_category、tb_defect_degree、tb_rating_level、tb_archive_item"],
            ["桥梁核心与档案", "tb_bridge、tb_bridge_basic_info、tb_bridge_component、tb_bridge_specific_component、tb_bridge_archive_record、tb_evaluation_history、tb_attachment"],
            ["初始检查", "tb_initial_inspection、tb_initial_inspection_item、tb_initial_component_inspection"],
            ["定期检查", "tb_periodic_inspection、tb_component_inspection、tb_defect"],
            ["任务与报告", "tb_inspection_task、tb_task_assignment、tb_task_status_history、tb_report"],
        ],
        [3.2, 11.3],
    )
    add_body(
        doc,
        "逻辑约束包括：桥梁编号、角色代码、登录账号和各类业务编码唯一；桥梁部件按桥梁编号和桥型部件配置编号组合唯一；桥梁具体部件按桥梁部件编号和部件序号组合唯一；桥梁基本信息按桥梁编号和版本号组合唯一，并保证一座桥梁仅一条当前有效版本；具体部件检查记录的桥梁具体部件编号非空。以上约束消除了原模型中从桥梁、配置、部位和标准部件直接连接具体部件的重复路径。",
    )
    add_picture(doc, LDM_DIAGRAM, 13.0)
    add_caption(doc, "图3.3 PowerDesigner逻辑数据模型（桥梁核心结构关系）")

    add_heading(doc, "3.2.3 物理模型设计", 3)
    add_body(
        doc,
        "目标物理模型选用MySQL 8.0、InnoDB和utf8mb4。关系修正以当前运行库为基线，但PDM表示后续迁移的目标结构：新增tb_bridge_component和tb_bridge_basic_info；tb_bridge_specific_component使用bridge_specific_component_id作为主键并通过bridge_component_id关联桥梁部件；tb_component_inspection通过bridge_specific_component_id关联被检查的具体部件。运行数据库应通过新的Flyway迁移逐步转换，不直接覆盖现有业务数据。",
    )
    add_table(
        doc,
        ["设计项", "物理设计"],
        [
            ["主键", "业务稳定对象使用VARCHAR编码主键，明细、病害、附件和日志使用BIGINT或INT自增主键"],
            ["外键", "32张业务表之间建立54条外键；桥梁核心默认图只显示经过校验的9条层级关系"],
            ["唯一约束", "桥型与标准部件、桥梁与桥型配置、桥梁部件与具体部件序号、桥梁与基本信息版本设置组合唯一约束"],
            ["索引", "桥梁路线、桥型、名称，桥梁部件父键、具体部件父键、检查日期、任务状态和日志时间等条件共建立69个索引"],
            ["状态与删除", "用户、桥梁、具体部件和业务记录使用状态字段或逻辑删除，归档检查与报告不执行级联物理删除"],
            ["文件存储", "照片、档案和报告保存于uploads受控目录，tb_attachment和tb_report保存相对路径、类型、大小和关联对象"],
            ["安全", "密码保存BCrypt散列，数据库凭据和JWT密钥由环境变量或部署配置提供，业务接口按角色鉴权"],
        ],
        [3.0, 11.5],
    )
    add_body(
        doc,
        "PDM中的九组核心引用均已绑定现有外键列，未产生PowerDesigner自动迁移的重复字段。部件表不再出现配置编号，桥梁类型部件配置表不再重复保存部位外键，桥梁具体部件表不再保存桥梁、桥型配置、部位和标准部件四条直接外键。数据库连接池使用HikariCP，涉及检查主表、具体部件检查、病害、附件和报告的组合操作必须在同一事务中提交或回滚。",
    )
    add_picture(doc, PDM_DIAGRAM, 13.0)
    add_caption(doc, "图3.4 PowerDesigner物理数据模型（桥梁核心结构关系）")

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    generate()
